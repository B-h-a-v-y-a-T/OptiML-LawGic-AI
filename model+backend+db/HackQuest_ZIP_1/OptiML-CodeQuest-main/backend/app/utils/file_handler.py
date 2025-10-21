import aiofiles
from pathlib import Path
from fastapi import UploadFile
import docx
import PyPDF2

TMP_DIR = Path("tmp")
TMP_DIR.mkdir(exist_ok=True)

async def extract_text_from_file(file: UploadFile) -> str:
    file_path = TMP_DIR / file.filename
    async with aiofiles.open(file_path, "wb") as out_file:
        content = await file.read()
        await out_file.write(content)

    text = ""
    if file.filename.endswith(".pdf"):
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join([page.extract_text() for page in reader.pages])
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])

    file_path.unlink()
    return text
import aiofiles
from pathlib import Path
from fastapi import UploadFile
import docx
import PyPDF2

TMP_DIR = Path("tmp")
TMP_DIR.mkdir(exist_ok=True)

async def extract_text_from_file(file: UploadFile) -> str:
    file_path = TMP_DIR / file.filename

    # Save uploaded file temporarily
    async with aiofiles.open(file_path, "wb") as out_file:
        content = await file.read()
        await out_file.write(content)

    text = ""
    if file.filename.endswith(".pdf"):
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join([page.extract_text() for page in reader.pages])
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])

    # Delete temporary file
    file_path.unlink()
    return text